"""Monthly roundup — Word export.

One .docx per (venue, month) roundup: the cross-committee state of play as
stored in the roundup's current summary version. Built entirely from
pipeline/briefing.py's editorial primitives (brand.py tokens) plus
docket_docx's generic markdown blocks, so the Word grammar matches the
meeting-briefing and docket exports — same fonts, eyebrows, takeaway rows.

Content comes from the DB only (current summary version, falling back to
the legacy report_md column); rendering is instant.
"""
from __future__ import annotations

import io
import re
from datetime import date

import pipeline.brand as brand
import pipeline.db as db
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from pipeline.briefing import (
    _CYAN,
    _CYAN_HEX,
    _GRAY_MID,
    _GRAY_MID_HEX,
    _GRAY_TEXT,
    _INK,
    _INK_SOFT,
    _LABEL,
    _eyebrow,
    _v2_page_number,
    _v2_pborder,
    _v2_right_tab,
    _v2_run,
    _v2_spacing,
)
from pipeline.docket_docx import _md_block, _split_h2, _takeaway_rows
from pipeline.roundup import month_label


def _as_date(d) -> date | None:
    if isinstance(d, date):
        return d
    try:
        return date.fromisoformat(str(d)[:10])
    except (TypeError, ValueError):
        return None


def _fmt_span(start, end) -> str:
    """'Mar 10' / 'Mar 10–11' / 'Mar 31–Apr 1' — compact source chips."""
    s = _as_date(start)
    if s is None:
        return ""
    e = _as_date(end)
    if e and e != s:
        if e.month == s.month:
            return f"{s.strftime('%b %-d')}–{e.strftime('%-d')}"
        return f"{s.strftime('%b %-d')}–{e.strftime('%b %-d')}"
    return s.strftime("%b %-d")


def generate_roundup_docx_bytes(roundup_id: int) -> tuple[bytes, str]:
    """Render the monthly roundup; returns (bytes, suggested_filename).
    Raises ValueError when the roundup or its report body is missing."""
    roundup = db.get_monthly_roundup(roundup_id)
    if not roundup:
        raise ValueError(f"Roundup {roundup_id} not found")

    report_md = (roundup.get("report_md") or "").strip()
    if not report_md:
        raise ValueError("This roundup has no report yet — generate it first")

    venue = roundup.get("venue_short") or ""
    label = month_label(roundup["month"])
    sources = db.get_roundup_meetings(roundup_id)
    today = date.today().strftime("%B %-d, %Y")

    doc = Document()

    # Page setup — render_briefing_docx's grammar.
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = brand.MARGIN_TOPBOT; sec.bottom_margin = brand.MARGIN_TOPBOT
    sec.left_margin = brand.MARGIN_SIDE; sec.right_margin = brand.MARGIN_SIDE
    content_w = (sec.page_width - sec.left_margin - sec.right_margin) // 635

    style = doc.styles["Normal"]
    style.font.name = brand.BODY_FONT; style.font.size = brand.SZ_BODY
    style.font.color.rgb = _INK_SOFT
    style.paragraph_format.line_spacing = brand.LINE_SPACING

    sec.different_first_page_header_footer = True
    if sec.first_page_header.paragraphs:
        sec.first_page_header.paragraphs[0].clear()
    if sec.first_page_footer.paragraphs:
        sec.first_page_footer.paragraphs[0].clear()
    hp = (sec.header.paragraphs[0] if sec.header.paragraphs
          else sec.header.add_paragraph())
    hp.clear()
    fp = (sec.footer.paragraphs[0] if sec.footer.paragraphs
          else sec.footer.add_paragraph())
    fp.clear()
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_c = OxmlElement("w:tab")
    tab_c.set(qn("w:val"), "center"); tab_c.set(qn("w:pos"), str(content_w // 2))
    tabs.append(tab_c)
    tab_r = OxmlElement("w:tab")
    tab_r.set(qn("w:val"), "right"); tab_r.set(qn("w:pos"), str(content_w))
    tabs.append(tab_r)
    pPr.append(tabs)
    _v2_run(fp, "Monthly Roundup", size=brand.SZ_FOOTER, color=_GRAY_TEXT,
            font=_LABEL)
    fp.add_run("\t")
    _v2_run(fp, "Page ", size=brand.SZ_FOOTER, color=_GRAY_TEXT, font=_LABEL)
    pr = fp.add_run()
    pr.font.name = _LABEL; pr.font.size = brand.SZ_FOOTER
    pr.font.color.rgb = _GRAY_TEXT
    _v2_page_number(pr)
    fp.add_run("\t")
    _v2_run(fp, f"{label} • {venue}", size=brand.SZ_FOOTER,
            color=_GRAY_TEXT, font=_LABEL)
    _v2_pborder(fp, "top", 6, _CYAN_HEX, space=4)

    if doc.paragraphs:
        doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

    # ── Cover header ────────────────────────────────────────────────────
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "top", 30, _CYAN_HEX)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(10), after=Pt(2))
    _v2_run(p, "POOLSIDE REPORTING SERVICE", size=brand.SZ_LABEL, bold=True,
            color=_CYAN, font=_LABEL, track=40)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(6))
    p.paragraph_format.keep_with_next = True
    _v2_run(p, label, size=brand.SZ_MASTHEAD, bold=True, color=_INK)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(6))
    _v2_right_tab(p, pos=content_w)
    _v2_run(p, f"{venue} Monthly Roundup", size=brand.SZ_HEADLINE, color=_CYAN,
            italic=True)
    p.add_run("\t")
    _v2_run(p, today, size=brand.SZ_HEADLINE, color=_GRAY_TEXT)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "bottom", 4, _GRAY_MID_HEX)

    # Provenance line, one hop under the rule: which briefings fed the month.
    if sources:
        chips = "  ·  ".join(
            f"{s.get('type_short', '')} {_fmt_span(s.get('meeting_date'), s.get('end_date'))}".strip()
            for s in sources
        )
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(9), after=Pt(0))
        _v2_run(p, f"Synthesized from {len(sources)} committee "
                   f"briefing{'s' if len(sources) != 1 else ''}:  ",
                size=brand.SZ_LINK, color=_GRAY_TEXT)
        _v2_run(p, chips, size=brand.SZ_LINK, color=_GRAY_MID, font=_LABEL)

    # ── Report body ─────────────────────────────────────────────────────
    preamble, sections = _split_h2(report_md)
    if preamble:
        _md_block(doc, preamble.splitlines())
    for heading, body in sections:
        _eyebrow(doc, heading.upper())
        is_takeaways = heading.strip().lower() == "key takeaways"
        bullets = [re.sub(r"^[-*]\s+", "", ln.strip())
                   for ln in body if re.match(r"^\s*[-*]\s+", ln)]
        if is_takeaways and bullets and all(
                not ln.strip() or re.match(r"^[-*]\s+|-{3,}", ln.strip())
                for ln in body):
            _takeaway_rows(doc, bullets)
        else:
            _md_block(doc, body)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    month_key = str(roundup["month"])[:7]
    filename = f"Roundup_{venue}_{month_key}.docx"
    return buf.read(), filename
