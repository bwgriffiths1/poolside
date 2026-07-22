"""
briefing.py — Word rendering for briefings and deep dives.

render_briefing_docx() walks the typed briefing AST produced by
api/briefing_parser.py (the same parse the web reader consumes) and renders
the NEPOOL-branded v2 design. Markdown is parsed exactly once, upstream —
this module deliberately contains NO briefing-markdown parser, so the Word
export cannot drift from the web view.

generate_deep_dive_docx_bytes() still parses deep-dive markdown locally
(deep dives have their own format and no web AST yet).

Inline images: <!-- image_id:N --> markers (and their web-resolved
![...](/api/images/N) form) are embedded as figures from document_images.
![...](/api/editor-images/N) — screenshots pasted into the editor — are
embedded from the separate editor_images table. Both kinds must be handled
here or they render as literal markdown text in the exported document.
"""
import base64
import logging
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, Twips, RGBColor
from docx.text.run import Run

from pipeline import brand

logger = logging.getLogger(__name__)

# Every form an image can take in briefing markdown: the summarizer's raw
# marker, the web-resolved document-figure link, and the editor's pasted
# screenshot link. Matching them in one pass keeps images in document order
# and guarantees each is consumed — an unmatched form would otherwise survive
# into the text runs as literal markdown, which is exactly how pasted
# screenshots used to export as a bare `![pasted](...)` stub.
_ANY_IMG_RE = re.compile(
    r"<!--\s*image_id:(?P<marker>\d+)\s*-->"
    r"|!\[[^\]]*\]\(/api/images/(?P<doc>\d+)\)"
    r"|!\[[^\]]*\]\(/api/editor-images/(?P<editor>\d+)\)"
)

# A figure spans the text column; the height cap keeps a tall screenshot from
# running off the page (pasted screenshots are far less predictably shaped
# than the figures pymupdf extracts from a PDF).
_MAX_IMG_W = 5.5   # inches
_MAX_IMG_H = 7.0   # inches

_MIME_SUFFIX = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/webp": ".webp",
}


def _fetch_image_record(image_id: int) -> dict | None:
    """Fetch an image record from DB by ID. Returns None on failure."""
    try:
        import pipeline.db as db
        rows = db.get_images_by_ids([image_id])
        return rows[0] if rows else None
    except Exception as exc:
        logger.warning("Failed to fetch image %d: %s", image_id, exc)
        return None


def _fetch_editor_image_record(image_id: int) -> dict | None:
    """Fetch an editor-pasted image record from DB by ID. None on failure."""
    try:
        import pipeline.db as db
        rows = db.get_editor_images_by_ids([image_id])
        return rows[0] if rows else None
    except Exception as exc:
        logger.warning("Failed to fetch editor image %d: %s", image_id, exc)
        return None


def _embed_image_bytes(
    doc: Document,
    img_bytes: bytes,
    *,
    caption: str = "",
    suffix: str = ".png",
    label: str = "image",
) -> None:
    """Embed raw image bytes as a figure, scaled to fit the text column.

    Shared by both image sources so extracted figures and pasted screenshots
    are sized and captioned identically.
    """
    # Write to temp file (python-docx needs a file path or file-like object)
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(img_bytes)
        tmp_path = tmp.name

    try:
        pic = doc.add_picture(tmp_path, width=Inches(_MAX_IMG_W))
        # add_picture preserves aspect ratio from the width alone; if that
        # leaves the figure taller than the page allows, re-fit by height.
        max_h = Inches(_MAX_IMG_H)
        if pic.height > max_h:
            pic.width = int(pic.width * (max_h / pic.height))
            pic.height = max_h
        # Centre the figure and give it air above.
        img_para = doc.paragraphs[-1]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _v2_spacing(img_para, before=Pt(8), after=Pt(0))
        # Keep the image on the same page as its caption below, if any.
        if caption:
            img_para.paragraph_format.keep_with_next = True
    except Exception as exc:
        logger.warning("Failed to embed %s in docx: %s", label, exc)
        return
    finally:
        try:
            Path(tmp_path).unlink()
        except OSError:
            pass

    if caption:
        # Editorial caption: mono, muted, centred, under a hairline — mirrors
        # the web .b-figure figcaption. Kept upper-case for the label feel.
        cap_para = doc.add_paragraph()
        _v2_spacing(cap_para, before=Pt(5), after=Pt(10))
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _v2_pborder(cap_para, "top", 4, _GRAY_MID_HEX, space=4)
        _v2_run(cap_para, caption.upper(), size=brand.SZ_CAPTION,
                color=_GRAY_TEXT, font=_LABEL)


def _add_image_to_doc(doc: Document, image_id: int) -> None:
    """Embed an extracted document figure (document_images) by DB id."""
    record = _fetch_image_record(image_id)
    if not record or not record.get("image_b64"):
        return
    try:
        img_bytes = base64.b64decode(record["image_b64"])
    except Exception as exc:
        logger.warning("Bad base64 for image %d: %s", image_id, exc)
        return
    _embed_image_bytes(
        doc,
        img_bytes,
        caption=record.get("description") or "",
        label=f"image {image_id}",
    )


def _add_editor_image_to_doc(doc: Document, image_id: int) -> None:
    """Embed an editor-pasted screenshot (editor_images) by DB id.

    These rows hold raw bytea rather than base64, and carry no description —
    so they render uncaptioned.
    """
    record = _fetch_editor_image_record(image_id)
    if not record or record.get("data") is None:
        return
    raw = record["data"]
    img_bytes = bytes(raw) if isinstance(raw, memoryview) else raw
    _embed_image_bytes(
        doc,
        img_bytes,
        suffix=_MIME_SUFFIX.get(record.get("mime_type") or "", ".png"),
        label=f"editor image {image_id}",
    )


def _add_matched_image(doc: Document, m: re.Match) -> None:
    """Dispatch one _ANY_IMG_RE match to the table that backs that form.

    Binds the immediately-preceding paragraph — typically the figure's
    "Figure: …" caption — to the image with keep-with-next, so a page break
    can't strand the caption at the foot of one page with its chart on the
    next.
    """
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.keep_with_next = True
    if m.group("editor"):
        _add_editor_image_to_doc(doc, int(m.group("editor")))
    else:
        _add_image_to_doc(doc, int(m.group("marker") or m.group("doc")))


# ---------------------------------------------------------------------------
# Date formatting helper
# ---------------------------------------------------------------------------

def _format_date_range(iso_dates: list[str]) -> str:
    """
    Format a list of ISO date strings as a human-readable range.

    ["2026-03-10"]                   → "March 10, 2026"
    ["2026-03-10", "2026-03-12"]     → "March 10–12, 2026"
    ["2026-03-31", "2026-04-01"]     → "March 31–April 1, 2026"
    """
    from datetime import date

    if not iso_dates:
        return ""
    dates = sorted(date.fromisoformat(d) for d in iso_dates)
    if len(dates) == 1:
        return dates[0].strftime("%B %-d, %Y")
    first, last = dates[0], dates[-1]
    if first.month == last.month and first.year == last.year:
        return f"{first.strftime('%B %-d')}–{last.strftime('%-d, %Y')}"
    if first.year == last.year:
        return f"{first.strftime('%B %-d')}–{last.strftime('%B %-d, %Y')}"
    return f"{first.strftime('%B %-d, %Y')}–{last.strftime('%B %-d, %Y')}"


# ---------------------------------------------------------------------------
# Redesigned briefing (v2) — NEPOOL brand design system
# ---------------------------------------------------------------------------

# Editorial palette — legacy constant names now resolve to the shared
# brand.py tokens (mirror of web/src/styles/tokens.css), so both this briefing
# renderer and the deep-dive path below share one palette. New code should
# prefer brand.* directly; these aliases keep the older structural code terse.
_INK        = brand.INK
_INK_SOFT   = brand.INK_SOFT
_CHARCOAL   = brand.INK          # was cyan-era charcoal; now editorial ink
_CYAN       = brand.ACCENT       # accent (terracotta)
_CYAN_BG    = brand.ACCENT_TINT  # tint band fill
_GRAY_BG    = brand.ELEV         # soft card fill
_GRAY_MID   = brand.MUTED_SOFT
_GRAY_TEXT  = brand.MUTED
_CYAN_HEX     = brand.ACCENT_HEX
_GRAY_MID_HEX = brand.BORDER     # hairline rule colour
_BORDER_SOFT  = brand.BORDER_SOFT
_BODY       = brand.BODY_FONT
_LABEL      = brand.LABEL_FONT
_CONTENT_W    = 9360  # 6.5" in twips


def _v2_run(para, text, *, size=brand.SZ_BODY, bold=False, color=_INK,
            italic=False, font=None):
    r = para.add_run(text)
    r.font.name = font or _BODY; r.font.size = size
    r.bold = bold; r.italic = italic; r.font.color.rgb = color
    return r


def _v2_link(para, url, text, *, size=brand.SZ_LINK, color=_CYAN, bold=False,
             underline=True, font=None):
    """Append a real, clickable external hyperlink run to `para`.

    python-docx has no public API for hyperlinks, so the relationship is
    registered on the part by hand and the run is wrapped in w:hyperlink.
    Formatting still goes through the normal Run wrapper.
    """
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    r_el = OxmlElement("w:r")
    link.append(r_el)
    para._p.append(link)

    run = Run(r_el, para)
    run.text = text
    run.font.name = font or _BODY; run.font.size = size
    run.bold = bold; run.font.color.rgb = color
    run.font.underline = underline
    return run


def _v2_pborder(para, side, sz, color, space=0):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr"); pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space)); el.set(qn("w:color"), color)
    pBdr.append(el)


def _v2_right_tab(para, pos=_CONTENT_W):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), str(pos))
    tabs.append(tab); pPr.append(tabs)


def _v2_cell_borders(cell, **sides):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side, attrs in sides.items():
        if attrs:
            el = OxmlElement(f"w:{side}")
            for k, v in attrs.items():
                el.set(qn(f"w:{k}"), str(v))
            tcB.append(el)
    tcPr.append(tcB)


def _v2_cell_shading(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill); tcPr.append(shd)


def _v2_cell_margins(cell, *, top=0, bottom=0, left=0, right=0):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _v2_pshading(para, fill):
    """Apply background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _v2_pindent(para, *, left=0, right=0):
    """Set paragraph left/right indent (twips)."""
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    if left: ind.set(qn("w:left"), str(left))
    if right: ind.set(qn("w:right"), str(right))
    pPr.append(ind)


def _v2_spacing(para, *, before=None, after=None, line=None):
    pf = para.paragraph_format
    if before is not None: pf.space_before = before
    if after is not None: pf.space_after = after
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line


# Inline markdown, matching web/src/lib/markdown.tsx::inline so the Word export
# renders exactly what the reader shows. Order matters: image before link.
_INLINE_ESC = set("$_*`[](){}#.-+!|<>%~&")
_INLINE_RE = re.compile(
    r"!\[([^\]]*)\]\(([^)\s]+)\)"     # 1 alt, 2 src   (image)
    r"|\[([^\]]+)\]\(([^)\s]+)\)"      # 3 text, 4 href (link)
    r"|\*\*([^*]+)\*\*"                # 5 bold
    r"|\*([^*]+)\*"                    # 6 italic
    r"|`([^`]+)`"                      # 7 code
)


def _unescape_inline(s: str) -> str:
    out: list[str] = []
    i = 0
    while i < len(s):
        if s[i] == "\\" and i + 1 < len(s) and s[i + 1] in _INLINE_ESC:
            out.append(s[i + 1]); i += 2
        else:
            out.append(s[i]); i += 1
    return "".join(out)


def _inline_runs(para, text, *, size=brand.SZ_BODY, color=_INK_SOFT, italic=False):
    """Render inline markdown into runs: **bold**, *italic*, `code`,
    [text](url). Bold lifts to ink; links become real clickable runs; code
    sets in the label mono. Backslash-escapes are unwound with the same
    character set the web uses so `\\$9,337/MWh` renders clean."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _v2_run(para, _unescape_inline(text[pos:m.start()]),
                    size=size, color=color, italic=italic)
        if m.group(2) is not None:            # image → fall back to alt text
            alt = _unescape_inline(m.group(1) or "")
            if alt:
                _v2_run(para, alt, size=size, color=color, italic=italic)
        elif m.group(4) is not None:          # link
            _v2_link(para, m.group(4), _unescape_inline(m.group(3)), size=size)
        elif m.group(5) is not None:          # bold
            _v2_run(para, _unescape_inline(m.group(5)),
                    size=size, bold=True, color=_INK, italic=italic)
        elif m.group(6) is not None:          # italic
            _v2_run(para, _unescape_inline(m.group(6)),
                    size=size, color=color, italic=True)
        elif m.group(7) is not None:          # code
            _v2_run(para, _unescape_inline(m.group(7)),
                    size=size, color=_INK, font=_LABEL)
        pos = m.end()
    if pos < len(text):
        _v2_run(para, _unescape_inline(text[pos:]), size=size, color=color, italic=italic)


def _v2_bold_runs(para, text, *, size=brand.SZ_BODY, color=_INK_SOFT):
    """Back-compat shim — all inline rendering now flows through _inline_runs."""
    _inline_runs(para, text, size=size, color=color)


def _smallcaps(run) -> None:
    """Set the small-caps run property (w:smallCaps) for editorial group heads."""
    rPr = run._r.get_or_add_rPr()
    el = OxmlElement("w:smallCaps")
    el.set(qn("w:val"), "1")
    rPr.append(el)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and len(s) > 2


def _is_separator_row(line: str) -> bool:
    """Return True for markdown separator lines like |---|---|"""
    s = line.strip()
    if not _is_table_row(s):
        return False
    return all(re.fullmatch(r"\s*:?-+:?\s*", cell) for cell in s[1:-1].split("|"))


def _parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip()[1:-1].split("|")]


_DELTA_POS = ("+",)
_DELTA_NEG = ("-", "−", "–")  # hyphen, minus sign, en-dash


def _add_word_table(doc: Document, rows: list[list[str]]) -> None:
    """Render cell-lists as a borderless editorial table (mirrors web .b-table):
    mono uppercase header under an accent rule, hairline row separators, no
    vertical lines; value columns right-aligned in the mono face with +/-
    deltas coloured. First row is the header."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.autofit = True
    for ri, row_data in enumerate(rows):
        is_header = ri == 0
        for ci in range(ncols):
            cell_text = (row_data[ci] if ci < len(row_data) else "").strip()
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            _v2_cell_margins(cell, top=54, bottom=54,
                             left=(0 if ci == 0 else 80), right=120)

            # Borders: borderless by default (header underline + row hairlines);
            # a full grid only when the shared TABLE_STYLE token asks for it.
            if brand.TABLE_STYLE == "ruled":
                edge = {"val": "single", "sz": 6, "color": _GRAY_MID_HEX, "space": 0}
                borders = {"top": edge, "bottom": edge, "left": edge, "right": edge}
            elif is_header:
                borders = {"bottom": {"val": "single", "sz": 12,
                                      "color": _CYAN_HEX, "space": 0}}
            else:
                borders = {"bottom": {"val": "single", "sz": 6,
                                      "color": _BORDER_SOFT, "space": 0}}
            _v2_cell_borders(cell, **borders)

            p = cell.paragraphs[0]
            _v2_spacing(p, before=Pt(0), after=Pt(0))
            value_col = ci > 0
            if value_col:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if is_header:
                _v2_run(p, cell_text.upper(), size=brand.SZ_LABEL, bold=True,
                        color=_GRAY_TEXT, font=_LABEL)
            elif value_col:
                delta = None
                if cell_text.startswith(_DELTA_POS):
                    delta = brand.SUCCESS
                elif cell_text.startswith(_DELTA_NEG):
                    delta = brand.DANGER
                _v2_run(p, cell_text, size=brand.SZ_BODY_SM,
                        color=delta or _INK, bold=bool(delta), font=_LABEL)
            else:
                _inline_runs(p, cell_text, size=brand.SZ_BODY_SM, color=_INK)


_CALLOUT_OPEN_RE = re.compile(r"^>\s*\[!([^\]]+)\]\s*(.*)$")


def _render_v2_callout(doc: Document, label: str, body_text: str) -> None:
    """Render a `> [!Label] body` admonition as a tinted left-bar callout box.

    Mirrors the web preview (.md-callout): cyan accent bar on the left,
    uppercase mono label, italic serif-style body. python-docx can't reproduce
    every nuance, so we approximate: Calibri italic body + bold caps label.
    """
    p = doc.add_paragraph()
    _v2_pshading(p, _CYAN_BG)
    _v2_pborder(p, "left", 24, _CYAN_HEX, space=8)
    _v2_pindent(p, left=180, right=120)
    _v2_spacing(p, before=Pt(8), after=Pt(8), line=1.3)
    p.paragraph_format.keep_together = True

    _v2_run(p, label.upper(), size=brand.SZ_LABEL, bold=True,
            color=_CYAN, font=_LABEL)
    p.add_run("\n")

    # Body — serif italic with inline bold/link support.
    body_text = body_text.strip()
    if body_text:
        _inline_runs(p, body_text, size=brand.SZ_BODY_SM, color=_INK, italic=True)


def _render_v2_body_lines(doc: Document, body_lines: list[str]) -> None:
    """Render body text lines, converting consecutive | rows into Word tables
    and <!-- image_id:N --> references into inline images."""
    i = 0
    while i < len(body_lines):
        line = body_lines[i]
        # Inline image reference — any of the three forms
        img_match = _ANY_IMG_RE.search(line)
        if img_match:
            _add_matched_image(doc, img_match)
            i += 1
            continue
        # Callout / admonition: `> [!Label] body…` + any continuation `>` lines.
        m_callout = _CALLOUT_OPEN_RE.match(line.strip())
        if m_callout:
            label = m_callout.group(1).strip()
            body_parts: list[str] = []
            if m_callout.group(2):
                body_parts.append(m_callout.group(2))
            i += 1
            while i < len(body_lines):
                cont = body_lines[i].lstrip()
                if not cont.startswith(">"):
                    break
                body_parts.append(re.sub(r"^>\s?", "", cont))
                i += 1
            _render_v2_callout(doc, label, " ".join(body_parts))
            continue
        # H4 sub-heading: italic blue number, bold italic black title
        if line.strip().startswith("#### "):
            h4_text = line.strip()[5:].strip()
            p = doc.add_paragraph()
            _v2_spacing(p, before=Pt(12), after=Pt(4))
            if ":" in h4_text:
                num, title = h4_text.split(":", 1)
                r = p.add_run(num.strip())
                r.font.name = "Calibri"; r.font.size = Pt(10.5)
                r.italic = True; r.font.color.rgb = _CYAN
                r = p.add_run(":  " + title.strip())
                r.font.name = "Calibri"; r.font.size = Pt(10.5)
                r.bold = True; r.italic = True; r.font.color.rgb = _CHARCOAL
            else:
                r = p.add_run(h4_text)
                r.font.name = "Calibri"; r.font.size = Pt(10.5)
                r.bold = True; r.italic = True; r.font.color.rgb = _CHARCOAL
            i += 1
            continue
        if _is_table_row(line):
            table_lines: list[str] = []
            while i < len(body_lines) and _is_table_row(body_lines[i]):
                table_lines.append(body_lines[i])
                i += 1
            data_rows = [_parse_table_row(r) for r in table_lines if not _is_separator_row(r)]
            _add_word_table(doc, data_rows)
            doc.add_paragraph()  # spacer after table
        else:
            p = doc.add_paragraph()
            _v2_spacing(p, before=Pt(0), after=Pt(7), line=Pt(12.1))
            _v2_bold_runs(p, line)
            i += 1


def _v2_page_number(run):
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
    run._r.extend([fc1, it, fc2])


def _render_v2_exec_summary(doc: Document, exec_lines: list[str]) -> None:
    """
    Render the executive summary as a single shaded box (like NEXT STEPS).
    All lines in one paragraph with \\n separators so the cyan background
    is one continuous block with no gaps.
    """
    if not exec_lines:
        return

    p = doc.add_paragraph()
    _v2_pshading(p, _CYAN_BG)
    _v2_pborder(p, "top", 10, _CYAN_HEX, space=6)
    _v2_pborder(p, "bottom", 10, _CYAN_HEX, space=6)
    _v2_spacing(p, before=Pt(0), after=Pt(0), line=Pt(12.1))

    for i, line in enumerate(exec_lines):
        if i > 0:
            p.add_run("\n")
        # Handle bullet lines
        if line.startswith("- ") or line.startswith("* "):
            _v2_run(p, "–  ", size=Pt(10.5), color=_CHARCOAL)
            _v2_bold_runs(p, line[2:])
        else:
            _v2_bold_runs(p, line)


# ---------------------------------------------------------------------------
# Briefing rendering — AST → NEPOOL-branded .docx bytes
# ---------------------------------------------------------------------------

def _render_v2_subheading(doc: Document, text: str) -> None:
    """Run-in sub-heading inside a section body (mirrors web .b-h3): serif and
    ink, with the leading number/label in accent when present."""
    p = doc.add_paragraph()
    _v2_spacing(p, before=Pt(12), after=Pt(4), line=brand.LINE_SPACING)
    p.paragraph_format.keep_with_next = True
    if ":" in text:
        num, title = text.split(":", 1)
        _v2_run(p, num.strip(), size=brand.SZ_H, bold=True, color=_CYAN)
        _v2_run(p, ":  " + title.strip(), size=brand.SZ_H, bold=True, color=_INK)
    else:
        _v2_run(p, text, size=brand.SZ_H, bold=True, color=_INK)


def _render_p_text(doc: Document, text: str) -> None:
    """Render plain paragraph text line by line, skipping blanks.
    Bullet lines arrive from the parser as `• ...`."""
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        if line.startswith("\u2022 "):
            _v2_spacing(p, before=Pt(0), after=Pt(3), line=brand.LINE_SPACING)
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.first_line_indent = Pt(-14)
            _v2_run(p, "\u2013  ", size=brand.SZ_BODY, bold=True, color=_CYAN)
            _inline_runs(p, line[2:], size=brand.SZ_BODY, color=_INK_SOFT)
        else:
            _v2_spacing(p, before=Pt(0), after=brand.SPACE_AFTER, line=brand.LINE_SPACING)
            _inline_runs(p, line, size=brand.SZ_BODY, color=_INK_SOFT)


def _render_p_block(doc: Document, text: str) -> None:
    """Render a paragraph block from the AST, embedding any image reference
    (summarizer marker, document figure, or pasted screenshot) at the point it
    appears so figures stay interleaved with the prose around them."""
    pos = 0
    for m in _ANY_IMG_RE.finditer(text):
        _render_p_text(doc, text[pos:m.start()])
        _add_matched_image(doc, m)
        pos = m.end()
    _render_p_text(doc, text[pos:])


def _render_section_docs(doc: Document, docs) -> None:
    """Materials for an agenda item as a grey left-bar list (same technique as
    the executive summary), NOT a table: a table's built-in cell inset pushes
    the box slightly off the body's left edge and is where Word's spacing goes
    haywire. Every paragraph shares one left indent, so the bar Word draws by
    merging the identical left borders is dead straight, and one filename per
    line (the extension is already in the name) keeps the left edge flush.

    Filenames with a scraped source_url become clickable (accent, no underline);
    the rest render as plain text so the file is still listed.
    """
    docs = list(docs or [])
    if not docs:
        return

    indent = Pt(12)  # constant for label + every row → the merged bar is straight

    def _bar(p, *, keep=True):
        _v2_pborder(p, "left", 16, brand.MUTED_SOFT_HEX, space=8)
        p.paragraph_format.left_indent = indent
        p.paragraph_format.widow_control = True
        if keep:
            p.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    _v2_spacing(p, before=Pt(2), after=Pt(3), line=brand.LINE_SPACING)
    _bar(p)
    _v2_run(p, "MATERIALS", size=brand.SZ_LABEL, bold=True, color=_GRAY_TEXT, font=_LABEL)

    for i, d in enumerate(docs):
        name = getattr(d, "filename", "") or ""
        url = getattr(d, "source_url", None)
        rp = doc.add_paragraph()
        _v2_spacing(rp, before=Pt(0), after=Pt(2), line=1.15)
        _bar(rp, keep=i < len(docs) - 1)
        if url:
            _v2_link(rp, url, name, size=brand.SZ_CAPTION, underline=False)
        else:
            _v2_run(rp, name, size=brand.SZ_CAPTION, color=_INK_SOFT)

    _v2_spacing(doc.add_paragraph(), before=Pt(0), after=Pt(4))


def _render_body_blocks(doc: Document, blocks) -> None:
    """Walk a section's typed body blocks (p / h / data / callout)."""
    for b in blocks or []:
        kind = getattr(b, "kind", "p")
        if kind == "data":
            rows = [list(r) for r in (getattr(b, "rows", None) or [])]
            if getattr(b, "title", ""):
                _render_v2_subheading(doc, b.title)
            if rows:
                _add_word_table(doc, rows)
                doc.add_paragraph()  # spacer after table
        elif kind == "callout":
            _render_v2_callout(doc, getattr(b, "label", "") or "Note",
                               getattr(b, "text", "") or "")
        elif kind == "h":
            _render_v2_subheading(doc, getattr(b, "text", "") or "")
            # Numbered sub-headings carry their own materials.
            _render_section_docs(doc, getattr(b, "docs", None))
        else:
            _render_p_block(doc, getattr(b, "text", "") or "")


def _exec_band(p) -> int:
    """Apply the executive-summary treatment to one paragraph and return the
    base left indent (points). Default 'band' runs a slim accent rule down the
    left of every exec paragraph \u2014 Word joins the per-paragraph left borders
    into one continuous bar; 'tint' shades the block; 'plain' adds nothing."""
    if brand.EXEC_TREATMENT == "band":
        _v2_pborder(p, "left", 18, _CYAN_HEX, space=9)
        return 11
    if brand.EXEC_TREATMENT == "tint":
        _v2_pshading(p, _CYAN_BG)
        p.paragraph_format.right_indent = Pt(6)
        return 8
    return 0


def _render_exec_blocks(doc: Document, blocks) -> None:
    """Executive summary as flowing prose (mirrors web .briefing-exec): each
    typed block rendered on its own \u2014 run-in bold sub-heads, spaced paragraphs,
    and real bullet lists \u2014 instead of one flat tint slab. The first prose
    paragraph reads as the lead (ink)."""
    if not blocks:
        return
    first_prose = True
    for b in blocks:
        kind = getattr(b, "kind", "p")
        text = (getattr(b, "text", "") or "").strip()
        if not text:
            continue

        if kind == "h":
            p = doc.add_paragraph()
            _v2_spacing(p, before=Pt(10), after=Pt(4), line=brand.LINE_SPACING)
            p.paragraph_format.keep_with_next = True
            base = _exec_band(p)
            if base:
                p.paragraph_format.left_indent = Pt(base)
            _v2_run(p, text, size=brand.SZ_H, bold=True, color=_INK)
            continue

        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        if lines and all(ln.startswith("\u2022 ") for ln in lines):
            for ln in lines:
                p = doc.add_paragraph()
                _v2_spacing(p, before=Pt(0), after=Pt(3), line=brand.LINE_SPACING)
                p.paragraph_format.widow_control = True
                base = _exec_band(p)
                p.paragraph_format.left_indent = Pt(base + 14)
                p.paragraph_format.first_line_indent = Pt(-14)
                _v2_run(p, "\u2013  ", size=brand.SZ_BODY, bold=True, color=_CYAN)
                _inline_runs(p, ln[2:], size=brand.SZ_BODY, color=_INK_SOFT)
        else:
            p = doc.add_paragraph()
            _v2_spacing(p, before=Pt(0), after=brand.SPACE_AFTER, line=brand.LINE_SPACING)
            p.paragraph_format.widow_control = True
            base = _exec_band(p)
            if base:
                p.paragraph_format.left_indent = Pt(base)
            _inline_runs(p, " ".join(lines), size=brand.SZ_BODY,
                         color=(_INK if first_prose else _INK_SOFT))
            first_prose = False


def _eyebrow(doc: Document, text: str):
    """Section eyebrow label (KEY TAKEAWAYS / EXECUTIVE SUMMARY / …): mono,
    accent, under an accent hairline. Mirrors the web .b-eyebrow."""
    p = doc.add_paragraph()
    _v2_spacing(p, before=Pt(22), after=Pt(10))
    p.paragraph_format.keep_with_next = True
    _v2_run(p, text, size=brand.SZ_EYEBROW, bold=True, color=_CYAN, font=_LABEL)
    _v2_pborder(p, "bottom", 10, _CYAN_HEX, space=4)
    return p


def render_briefing_docx(
    briefing,
    committee: str,
    meeting_dates: list[str],
    materials_url: str | None = None,
    webex_url: str | None = None,
) -> bytes:
    """
    Render a parsed briefing AST to the NEPOOL-branded v2 .docx design and
    return raw bytes.

    `briefing` is the object api/briefing_parser.py produces for the web
    reader (duck-typed: .tldr, .executive_summary, .sections; sections carry
    .item_id/.depth/.title/.body/.next_steps; blocks carry .kind and
    kind-specific fields). Consuming the SAME parse as the web page is the
    point: markdown is parsed exactly once, so the Word export can no longer
    silently drift from what the reader shows (which happened twice while
    there was a second parser here).

    `materials_url` / `webex_url` are the venue-hosted links (see
    pipeline/venue_links.py); each renders under the cover header when given.
    """
    import io

    date_str = _format_date_range(meeting_dates)

    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = brand.MARGIN_TOPBOT; sec.bottom_margin = brand.MARGIN_TOPBOT
    sec.left_margin = brand.MARGIN_SIDE; sec.right_margin = brand.MARGIN_SIDE
    content_w = (sec.page_width - sec.left_margin - sec.right_margin) // 635  # twips

    style = doc.styles["Normal"]
    style.font.name = _BODY; style.font.size = brand.SZ_BODY
    style.font.color.rgb = _INK_SOFT
    style.paragraph_format.line_spacing = brand.LINE_SPACING

    sec.different_first_page_header_footer = True
    if sec.first_page_header.paragraphs:
        sec.first_page_header.paragraphs[0].clear()
    if sec.first_page_footer.paragraphs:
        sec.first_page_footer.paragraphs[0].clear()
    # Running header — empty
    hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
    hp.clear()
    # Running footer — LEFT: "Meeting Briefing" | CENTER: Page N | RIGHT: Date • Committee
    fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.clear()
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_c = OxmlElement("w:tab")
    tab_c.set(qn("w:val"), "center"); tab_c.set(qn("w:pos"), str(content_w // 2))
    tabs.append(tab_c)
    tab_r = OxmlElement("w:tab")
    tab_r.set(qn("w:val"), "right"); tab_r.set(qn("w:pos"), str(content_w))
    tabs.append(tab_r)
    pPr.append(tabs)
    _v2_run(fp, "Meeting Briefing", size=brand.SZ_FOOTER, color=_GRAY_TEXT, font=_LABEL)
    fp.add_run("\t")
    _v2_run(fp, "Page ", size=brand.SZ_FOOTER, color=_GRAY_TEXT, font=_LABEL)
    pr = fp.add_run(); pr.font.name = _LABEL; pr.font.size = brand.SZ_FOOTER; pr.font.color.rgb = _GRAY_TEXT
    _v2_page_number(pr)
    fp.add_run("\t")
    _v2_run(fp, f"{date_str} \u2022 {committee}", size=brand.SZ_FOOTER, color=_GRAY_TEXT, font=_LABEL)
    _v2_pborder(fp, "top", 6, _CYAN_HEX, space=4)

    if doc.paragraphs:
        doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "top", 30, _CYAN_HEX)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(10), after=Pt(2))
    _v2_run(p, "N E P O O L", size=brand.SZ_LABEL, bold=True, color=_CYAN, font=_LABEL)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(6))
    p.paragraph_format.keep_with_next = True
    _v2_run(p, committee, size=brand.SZ_MASTHEAD, bold=True, color=_INK)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(12)); _v2_right_tab(p, pos=content_w)
    _v2_run(p, "Meeting Briefing", size=brand.SZ_HEADLINE, color=_CYAN, italic=True)
    p.add_run("\t")
    _v2_run(p, date_str, size=brand.SZ_HEADLINE, color=_GRAY_TEXT)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "bottom", 4, _GRAY_MID_HEX)

    # Venue links, directly under the header rule — the source materials and
    # the virtual-attendance permalink, one hop from the briefing.
    if materials_url or webex_url:
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(9), after=Pt(0))
        if materials_url:
            _v2_run(p, "Meeting materials:  ", size=brand.SZ_LINK, color=_GRAY_TEXT)
            _v2_link(p, materials_url, "View on iso-ne.com")
        if webex_url:
            if materials_url:
                _v2_run(p, "      •      ", size=brand.SZ_LINK, color=_GRAY_MID)
            _v2_run(p, "Join virtually:  ", size=brand.SZ_LINK, color=_GRAY_TEXT)
            _v2_link(p, webex_url, "ISO-NE Webex")

    tldr = list(getattr(briefing, "tldr", None) or [])
    if tldr:
        _eyebrow(doc, "KEY TAKEAWAYS")
        gutter = Pt(22)
        for rank, tk in enumerate(tldr, 1):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            _v2_spacing(p, before=Pt(4), after=Pt(4), line=brand.LINE_SPACING)
            # Hanging indent + a tab stop at the gutter: the number sits in a
            # fixed-width gutter and a real tab (not spaces) advances the text to
            # the gutter, so the body's left edge is flush on the first line and
            # every wrapped line — no raggedness from the number's width.
            pf.left_indent = gutter
            pf.first_line_indent = -gutter
            pf.tab_stops.add_tab_stop(gutter)
            pf.widow_control = True
            # `w:between` (not a per-paragraph top border) — Word collapses
            # identical top/bottom borders on consecutive paragraphs into one
            # outer rule, so top borders drew a line only above item 2. A
            # `between` border on every (identically-bordered) takeaway is
            # rendered between each consecutive pair, giving a hairline between
            # all of them.
            _v2_pborder(p, "between", 4, _BORDER_SOFT, space=6)
            _v2_run(p, f"{rank}.", size=brand.SZ_BODY, bold=True, color=_CYAN, font=_LABEL)
            p.add_run()._r.append(OxmlElement("w:tab"))
            _inline_runs(p, tk, size=brand.SZ_BODY, color=_INK)

    exec_blocks = list(getattr(briefing, "executive_summary", None) or [])
    if exec_blocks:
        _eyebrow(doc, "EXECUTIVE SUMMARY")
        _render_exec_blocks(doc, exec_blocks)

    sections = list(getattr(briefing, "sections", None) or [])
    if sections:
        _eyebrow(doc, "AGENDA ITEM SUMMARIES")

    for item in sections:
        depth = getattr(item, "depth", 0)
        number = getattr(item, "item_id", "") or ""
        title = getattr(item, "title", "") or ""
        if depth == 0:
            # Top-level agenda item — small-caps, rule-underlined group header.
            p = doc.add_paragraph(); _v2_spacing(p, before=Pt(22), after=Pt(7))
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            _v2_run(p, number, size=brand.SZ_GROUP, bold=True, color=_CYAN)
            tr = _v2_run(p, "  " + title, size=brand.SZ_GROUP, bold=True, color=_INK)
            if brand.SMALL_CAPS_GROUPS:
                _smallcaps(tr)
            _v2_pborder(p, "bottom", 8, _GRAY_MID_HEX, space=3)
        else:
            p = doc.add_paragraph(); _v2_spacing(p, before=Pt(16), after=Pt(6))
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            p.paragraph_format.left_indent = Pt(14)
            _v2_run(p, number, size=brand.SZ_SUBITEM, bold=True, color=_CYAN)
            _v2_run(p, "  " + title, size=brand.SZ_SUBITEM, bold=True, color=_INK)

        _render_section_docs(doc, getattr(item, "docs", None))
        _render_body_blocks(doc, getattr(item, "body", None))

        next_steps = getattr(item, "next_steps", None) or []
        if next_steps:
            doc.add_paragraph()
            p = doc.add_paragraph()
            _v2_pshading(p, _GRAY_BG)
            _v2_pborder(p, "top", 8, _GRAY_MID_HEX, space=6)
            _v2_pborder(p, "bottom", 8, _GRAY_MID_HEX, space=6)
            _v2_pindent(p, left=140, right=120)
            _v2_spacing(p, before=Pt(0), after=Pt(0), line=brand.LINE_SPACING)
            p.paragraph_format.keep_together = True
            _v2_run(p, "NEXT STEPS", size=brand.SZ_LABEL, bold=True, color=_GRAY_TEXT, font=_LABEL)
            for step in next_steps:
                p.add_run("\n")
                _v2_run(p, "\u2192  ", size=brand.SZ_BODY_SM, bold=True, color=_CYAN, font=_LABEL)
                _inline_runs(p, step, size=brand.SZ_BODY_SM, color=_INK)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Deep dive / special report rendering
# ---------------------------------------------------------------------------

def _parse_deep_dive_md(text: str) -> dict:
    """
    Parse deep-dive report markdown into structured data.
    Returns {"title", "date", "sections": [{"heading", "body": [str]}]}.
    Generic section parser — not hardcoded to briefing structure.
    """
    lines = text.splitlines()
    data: dict = {"title": "", "date": "", "sections": []}
    i = 0

    # Optional H1 title
    for j, line in enumerate(lines):
        s = line.strip()
        if s.startswith("## "):
            break
        if s.startswith("# "):
            data["title"] = s[2:].strip()
            i = j + 1
            break

    # Scan for ## sections
    current_section: dict | None = None
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("## "):
            if current_section:
                data["sections"].append(current_section)
            heading = s[3:].strip()
            current_section = {"heading": heading, "body": []}
            i += 1
            continue
        if s == "---":
            i += 1
            continue
        if current_section is not None and s:
            current_section["body"].append(s)
        elif current_section is not None and not s:
            # Preserve paragraph breaks as empty strings
            current_section["body"].append("")
        i += 1

    if current_section:
        data["sections"].append(current_section)

    return data


def generate_deep_dive_docx_bytes(
    report_md: str,
    title: str,
    document_names: list[str],
    date_range: str,
) -> bytes:
    """
    Render a deep dive report as a .docx using the NEPOOL brand design.
    Returns raw .docx bytes.
    """
    import io

    data = _parse_deep_dive_md(report_md)
    data["title"] = data["title"] or title
    data["date"] = data["date"] or date_range

    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Twips(1008); sec.bottom_margin = Twips(1008)
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(10.5)
    style.font.color.rgb = _CHARCOAL

    # Header / footer
    sec.different_first_page_header_footer = True
    if sec.first_page_header.paragraphs:
        sec.first_page_header.paragraphs[0].clear()
    if sec.first_page_footer.paragraphs:
        sec.first_page_footer.paragraphs[0].clear()
    hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
    hp.clear()
    # Running footer — LEFT: "Special Report" | CENTER: Page N | RIGHT: Date • Title
    fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.clear()
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_c = OxmlElement("w:tab")
    tab_c.set(qn("w:val"), "center"); tab_c.set(qn("w:pos"), str(_CONTENT_W // 2))
    tabs.append(tab_c)
    tab_r = OxmlElement("w:tab")
    tab_r.set(qn("w:val"), "right"); tab_r.set(qn("w:pos"), str(_CONTENT_W))
    tabs.append(tab_r)
    pPr.append(tabs)
    _v2_run(fp, "Special Report", size=Pt(8.5), color=_GRAY_MID)
    fp.add_run("\t")
    _v2_run(fp, "Page ", size=Pt(8.5), color=_GRAY_MID)
    pr = fp.add_run(); pr.font.name = "Calibri"; pr.font.size = Pt(8.5); pr.font.color.rgb = _GRAY_MID
    _v2_page_number(pr)
    fp.add_run("\t")
    _v2_run(fp, f"{data['date']} • {data['title']}", size=Pt(8.5), color=_GRAY_MID)
    _v2_pborder(fp, "top", 6, _CYAN_HEX, space=4)

    # Remove default empty paragraph
    if doc.paragraphs:
        doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

    # ---- Title block ----
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "top", 36, _CYAN_HEX)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_run(p, "N E P O O L", size=Pt(8), bold=True, color=_CYAN)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(4))
    _v2_run(p, data["title"], size=Pt(28), bold=True, color=_CHARCOAL)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(10)); _v2_right_tab(p)
    _v2_run(p, "Special Report", size=Pt(12), color=_CYAN)
    p.add_run("\t")
    _v2_run(p, data["date"], size=Pt(12), color=_GRAY_TEXT)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "bottom", 4, _GRAY_MID_HEX)

    # ---- Source documents listing ----
    if document_names:
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(10), after=Pt(6))
        _v2_run(p, "Source documents: ", size=Pt(9), bold=True, color=_CHARCOAL)
        _v2_run(p, ", ".join(document_names), size=Pt(9), color=_GRAY_TEXT)

    # ---- Sections ----
    for section in data["sections"]:
        heading = section["heading"]
        body = section["body"]

        # Section heading — cyan-underlined like EXECUTIVE SUMMARY / AGENDA ITEM SUMMARIES
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(22), after=Pt(11))
        _v2_run(p, heading.upper(), size=Pt(10), bold=True, color=_CHARCOAL)
        _v2_pborder(p, "bottom", 8, _CYAN_HEX, space=4)

        # Executive Summary gets the shaded box treatment
        if "executive summary" in heading.lower():
            _render_v2_exec_summary(doc, [l for l in body if l])
        else:
            _render_v2_body_lines(doc, body)

    # ---- Save to bytes ----
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
