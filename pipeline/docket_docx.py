"""FERC docket briefing — Word export.

One .docx per docket: the state of play up front (Key Takeaways in the
briefing band treatment), then one page per substantive filing with its
eLibrary links and file list. Built entirely from pipeline/briefing.py's
editorial primitives (brand.py tokens) so the Word grammar matches the
meeting briefing export — same fonts, eyebrows, grey-bar materials.

Content comes from the DB only (SOP current version + filing summaries);
no FERC calls, so rendering is instant.
"""
from __future__ import annotations

import io
import re
from datetime import date

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import pipeline.brand as brand
import pipeline.db as db
from pipeline.briefing import (
    _CYAN,
    _CYAN_HEX,
    _GRAY_MID,
    _GRAY_MID_HEX,
    _GRAY_TEXT,
    _INK,
    _INK_SOFT,
    _LABEL,
    _eyebrow,
    _inline_runs,
    _render_v2_body_lines,
    _render_v2_subheading,
    _v2_link,
    _v2_page_number,
    _v2_pborder,
    _v2_right_tab,
    _v2_run,
    _v2_spacing,
)
from pipeline.docket_ingest import author_orgs
from pipeline.ferc_client import docinfo_url, filelist_url

# Web classLabel's mapping, kept small on purpose.
_CLASS_LABEL = {
    "Application/Petition/Request": "Filing",
    "Comments/Protest": "Comments",
    "Order/Opinion": "Order",
    "ALJ Issuance": "ALJ Issuance",
    "Pleading/Motion": "Motion",
    "Briefing/Arguments of Law": "Brief",
    "Testimony": "Testimony",
}


def _class_label(f: dict) -> str:
    if f.get("role") == "initial":
        return "Initial Filing"
    return _CLASS_LABEL.get(f.get("document_class") or "", f.get("document_class") or "Filing")


def _fmt_date(d) -> str:
    if not d:
        return ""
    if isinstance(d, str):
        try:
            d = date.fromisoformat(d[:10])
        except ValueError:
            return d
    return d.strftime("%B %-d, %Y")


def _fmt_size(n) -> str:
    if not n:
        return ""
    if n > 1_000_000:
        return f"{n / 1_000_000:.1f} MB"
    return f"{round(n / 1000)} KB"


# ── markdown → docx ─────────────────────────────────────────────────────

def _split_h2(md: str) -> tuple[str, list[tuple[str, list[str]]]]:
    """(preamble, [(heading, body_lines)]) split at `## ` lines."""
    preamble: list[str] = []
    sections: list[tuple[str, list[str]]] = []
    cur: tuple[str, list[str]] | None = None
    for line in (md or "").splitlines():
        m = re.match(r"^##\s+(.+?)\s*$", line)
        if m:
            if cur:
                sections.append(cur)
            cur = (m.group(1).strip(), [])
        elif cur:
            cur[1].append(line)
        else:
            preamble.append(line)
    if cur:
        sections.append(cur)
    return "\n".join(preamble).strip(), sections


def _md_block(doc: Document, lines: list[str]) -> None:
    """Generic markdown body: ### subheads and - bullets handled here,
    everything else (tables, callouts, bold runs) delegated to briefing's
    line renderer in buffered chunks."""
    buf: list[str] = []

    def flush() -> None:
        if buf:
            _render_v2_body_lines(doc, [ln for ln in buf if ln.strip()])
            buf.clear()

    for raw in lines:
        line = raw.rstrip()
        s = line.strip()
        if not s or re.fullmatch(r"-{3,}", s):
            continue
        if s.startswith("### "):
            flush()
            _render_v2_subheading(doc, s[4:].strip())
        elif s.startswith(">") and not s.startswith("> [!"):
            # Bare blockquote (analyst notes): accent-barred italic para.
            # `> [!Label]` callouts fall through to the line renderer.
            flush()
            p = doc.add_paragraph()
            _v2_spacing(p, before=Pt(4), after=Pt(8), line=brand.LINE_SPACING)
            _v2_pborder(p, "left", 12, _CYAN_HEX, space=8)
            p.paragraph_format.left_indent = Pt(10)
            _inline_runs(p, re.sub(r"^>\s?", "", s), size=brand.SZ_BODY,
                         color=_INK_SOFT, italic=True)
        elif re.match(r"^[-*]\s+", s):
            flush()
            p = doc.add_paragraph()
            _v2_spacing(p, before=Pt(0), after=Pt(3), line=brand.LINE_SPACING)
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.first_line_indent = Pt(-14)
            p.paragraph_format.widow_control = True
            _v2_run(p, "–  ", size=brand.SZ_BODY, bold=True, color=_CYAN)
            _inline_runs(p, re.sub(r"^[-*]\s+", "", s),
                         size=brand.SZ_BODY, color=_INK_SOFT)
        else:
            buf.append(s)
    flush()


def _takeaway_rows(doc: Document, bullets: list[str]) -> None:
    """The briefing's KEY TAKEAWAYS treatment: numbered gutter rows with a
    `w:between` rule (same Word gotcha notes as render_briefing_docx)."""
    gutter = Pt(22)
    for rank, tk in enumerate(bullets, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        _v2_spacing(p, before=Pt(4), after=Pt(4), line=brand.LINE_SPACING)
        pf.left_indent = gutter
        pf.first_line_indent = -gutter
        pf.tab_stops.add_tab_stop(gutter)
        pf.widow_control = True
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        el = OxmlElement("w:between")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "3")
        el.set(qn("w:color"), brand.BORDER_SOFT)
        pBdr.append(el)
        pPr.append(pBdr)
        _v2_run(p, f"{rank:02d}", size=Pt(10), bold=True, color=_GRAY_MID,
                font=_LABEL)
        p.add_run("\t")
        _inline_runs(p, tk, size=brand.SZ_BODY, color=_INK)


def _grey_bar_files(doc: Document, files: list[dict]) -> None:
    """The materials grey-bar block (briefing._render_section_docs shape):
    one file per line with its size; excluded tariff sheets annotated."""
    files = [x for x in files or []]
    if not files:
        return
    indent = Pt(12)

    def _bar(p, *, keep=True):
        _v2_pborder(p, "left", 16, brand.MUTED_SOFT_HEX, space=8)
        p.paragraph_format.left_indent = indent
        p.paragraph_format.widow_control = True
        if keep:
            p.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    _v2_spacing(p, before=Pt(2), after=Pt(3), line=brand.LINE_SPACING)
    _bar(p)
    _v2_run(p, "FILES", size=brand.SZ_LABEL, bold=True, color=_GRAY_TEXT,
            font=_LABEL)

    for i, x in enumerate(files):
        name = x.get("file_desc") or x.get("orig_file_name") or "file"
        meta_bits = []
        if x.get("page_count") and x["page_count"] > 1:
            meta_bits.append(f"{x['page_count']}pp")
        if x.get("file_size"):
            meta_bits.append(_fmt_size(x["file_size"]))
        if not x.get("included"):
            meta_bits.append("not summarized")
        rp = doc.add_paragraph()
        _v2_spacing(rp, before=Pt(0), after=Pt(2), line=1.15)
        _bar(rp, keep=i < len(files) - 1)
        _v2_run(rp, name, size=brand.SZ_CAPTION,
                color=(_INK_SOFT if x.get("included") else _GRAY_MID))
        if meta_bits:
            _v2_run(rp, f"  —  {' · '.join(meta_bits)}",
                    size=brand.SZ_CAPTION, color=_GRAY_MID, font=_LABEL)

    _v2_spacing(doc.add_paragraph(), before=Pt(0), after=Pt(4))


# ── document assembly ───────────────────────────────────────────────────

def generate_docket_docx_bytes(docket_id: int) -> tuple[bytes, str]:
    """Render the docket briefing; returns (bytes, suggested_filename).
    Raises ValueError when the docket or its state of play is missing."""
    docket = db.get_docket(docket_id)
    if not docket:
        raise ValueError(f"Docket {docket_id} not found")
    number = docket["docket_number"]

    sop_row = db.get_current_summary("docket", docket_id)
    sop_md = (sop_row or {}).get("detailed") or ""
    if not sop_md.strip():
        raise ValueError(f"No state of play on {number} yet — run a sync first")

    filings = db.list_docket_filings(docket_id)
    files_by_filing: dict[int, list[dict]] = {}
    for x in db.list_docket_filing_files(docket_id):
        files_by_filing.setdefault(x["filing_id"], []).append(x)

    # Substantive filings, oldest first — the document reads as a record.
    # Treatment decides (not class): a doc-ful Intervention carrying a
    # protest gets its page like any other substantive filing.
    subs = sorted(
        [f for f in filings if f.get("treatment") != "skip"],
        key=lambda f: str(f.get("filed_date") or f.get("issued_date") or ""),
    )
    intervenors = sorted(
        {org for f in filings if f.get("document_class") == "Intervention"
         for org in author_orgs(f.get("filing_parties"))},
        key=str.casefold,
    )

    today = date.today().strftime("%B %-d, %Y")

    doc = Document()

    # Page setup — render_briefing_docx's grammar.
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = brand.MARGIN_TOPBOT; sec.bottom_margin = brand.MARGIN_TOPBOT
    sec.left_margin = brand.MARGIN_SIDE; sec.right_margin = brand.MARGIN_SIDE
    content_w = (sec.page_width - sec.left_margin - sec.right_margin) // 635

    style = doc.styles["Normal"]
    style.font.name = brand.BODY_FONT; style.font.size = brand.SZ_BODY
    style.font.color.rgb = _INK_SOFT
    style.paragraph_format.line_spacing = brand.LINE_SPACING

    sec.different_first_page_header_footer = True
    if sec.first_page_header.paragraphs:
        sec.first_page_header.paragraphs[0].clear()
    if sec.first_page_footer.paragraphs:
        sec.first_page_footer.paragraphs[0].clear()
    hp = (sec.header.paragraphs[0] if sec.header.paragraphs
          else sec.header.add_paragraph())
    hp.clear()
    fp = (sec.footer.paragraphs[0] if sec.footer.paragraphs
          else sec.footer.add_paragraph())
    fp.clear()
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_c = OxmlElement("w:tab")
    tab_c.set(qn("w:val"), "center"); tab_c.set(qn("w:pos"), str(content_w // 2))
    tabs.append(tab_c)
    tab_r = OxmlElement("w:tab")
    tab_r.set(qn("w:val"), "right"); tab_r.set(qn("w:pos"), str(content_w))
    tabs.append(tab_r)
    pPr.append(tabs)
    _v2_run(fp, "Docket Briefing", size=brand.SZ_FOOTER, color=_GRAY_TEXT,
            font=_LABEL)
    fp.add_run("\t")
    _v2_run(fp, "Page ", size=brand.SZ_FOOTER, color=_GRAY_TEXT, font=_LABEL)
    pr = fp.add_run()
    pr.font.name = _LABEL; pr.font.size = brand.SZ_FOOTER
    pr.font.color.rgb = _GRAY_TEXT
    _v2_page_number(pr)
    fp.add_run("\t")
    _v2_run(fp, f"{today} • {number}", size=brand.SZ_FOOTER,
            color=_GRAY_TEXT, font=_LABEL)
    _v2_pborder(fp, "top", 6, _CYAN_HEX, space=4)

    if doc.paragraphs:
        doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

    # ── Cover header ────────────────────────────────────────────────────
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "top", 30, _CYAN_HEX)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(10), after=Pt(2))
    _v2_run(p, "P O O L S I D E", size=brand.SZ_LABEL, bold=True, color=_CYAN,
            font=_LABEL)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(6))
    p.paragraph_format.keep_with_next = True
    _v2_run(p, number, size=brand.SZ_MASTHEAD, bold=True, color=_INK)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(6))
    _v2_right_tab(p, pos=content_w)
    _v2_run(p, "Docket State of Play", size=brand.SZ_HEADLINE, color=_CYAN,
            italic=True)
    p.add_run("\t")
    _v2_run(p, today, size=brand.SZ_HEADLINE, color=_GRAY_TEXT)
    if docket.get("title"):
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(10))
        _v2_run(p, docket["title"], size=Pt(11), color=_GRAY_TEXT, italic=True)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "bottom", 4, _GRAY_MID_HEX)

    # Docket record link + counts, one hop under the rule.
    root = next((f for f in subs
                 if f.get("document_class") == "Application/Petition/Request"),
                subs[0] if subs else None)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(9), after=Pt(0))
    _v2_run(p, f"{len(filings)} filings • {len(intervenors)} intervenors",
            size=brand.SZ_LINK, color=_GRAY_TEXT)
    if root:
        _v2_run(p, "      •      ", size=brand.SZ_LINK, color=_GRAY_MID)
        _v2_run(p, "Docket record:  ", size=brand.SZ_LINK, color=_GRAY_TEXT)
        _v2_link(p, docinfo_url(root["accession_number"]), "FERC eLibrary")

    # ── State of play ───────────────────────────────────────────────────
    preamble, sop_sections = _split_h2(sop_md)
    if preamble:
        _md_block(doc, preamble.splitlines())
    for heading, body in sop_sections:
        _eyebrow(doc, heading.upper())
        is_takeaways = heading.strip().lower() == "key takeaways"
        bullets = [re.sub(r"^[-*]\s+", "", ln.strip())
                   for ln in body if re.match(r"^\s*[-*]\s+", ln)]
        if is_takeaways and bullets and all(
                not ln.strip() or re.match(r"^[-*]\s+|-{3,}", ln.strip())
                for ln in body):
            _takeaway_rows(doc, bullets)
        else:
            _md_block(doc, body)

    # ── Intervenors ─────────────────────────────────────────────────────
    if intervenors:
        _eyebrow(doc, f"INTERVENORS ({len(intervenors)})")
        p = doc.add_paragraph()
        _v2_spacing(p, before=Pt(0), after=Pt(6), line=brand.LINE_SPACING)
        _v2_run(p, "  •  ".join(intervenors), size=brand.SZ_BODY_SM,
                color=_INK_SOFT)

    # ── One page per substantive filing ─────────────────────────────────
    total = len(subs)
    for i, f in enumerate(subs, 1):
        pb = doc.add_paragraph()
        pb.add_run().add_break(WD_BREAK.PAGE)

        label = _class_label(f)
        who = "; ".join(author_orgs(f.get("filing_parties"))) or "Unknown party"

        eyebrow = f"FILING {i} OF {total}"
        if f.get("role") == "initial":
            eyebrow += "  —  THE INITIAL FILING"
        elif f.get("role") == "order":
            eyebrow += "  —  FERC ORDER"
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(2))
        _v2_run(p, eyebrow, size=brand.SZ_EYEBROW, bold=True,
                color=_CYAN, font=_LABEL)
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(4))
        p.paragraph_format.keep_with_next = True
        _v2_run(p, f"{label}: {who}", size=brand.SZ_GROUP, bold=True,
                color=_INK)
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(6))
        _v2_right_tab(p, pos=content_w)
        _v2_run(p, _fmt_date(f.get("filed_date") or f.get("issued_date")),
                size=Pt(11), color=_GRAY_TEXT)
        p.add_run("\t")
        _v2_run(p, f.get("accession_number") or "", size=Pt(10),
                color=_GRAY_MID, font=_LABEL)
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
        _v2_pborder(p, "bottom", 4, _GRAY_MID_HEX)

        # Meta + source links.
        meta_bits = []
        if f.get("ferc_cite"):
            meta_bits.append(f"Cite {f['ferc_cite']}")
        if f.get("comments_due_date"):
            meta_bits.append(f"Comments due {_fmt_date(f['comments_due_date'])}")
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(8), after=Pt(4))
        if meta_bits:
            _v2_run(p, "  •  ".join(meta_bits) + "      ",
                    size=brand.SZ_LINK, color=_GRAY_TEXT)
        _v2_run(p, "Source:  ", size=brand.SZ_LINK, color=_GRAY_TEXT)
        _v2_link(p, docinfo_url(f["accession_number"]), "Doc info")
        _v2_run(p, "  ·  ", size=brand.SZ_LINK, color=_GRAY_MID)
        _v2_link(p, filelist_url(f["accession_number"]), "File list")

        if f.get("description"):
            p = doc.add_paragraph(); _v2_spacing(p, before=Pt(2), after=Pt(8))
            _v2_run(p, f["description"], size=brand.SZ_BODY_SM,
                    color=_GRAY_TEXT, italic=True)

        _grey_bar_files(doc, files_by_filing.get(f["id"], []))

        summary = f.get("summary_detailed") or ""
        if f.get("summary_one_line"):
            p = doc.add_paragraph()
            _v2_spacing(p, before=Pt(4), after=Pt(8), line=brand.LINE_SPACING)
            _inline_runs(p, f["summary_one_line"], size=brand.SZ_BODY,
                         color=_INK)
        if summary.strip():
            s_pre, s_sections = _split_h2(summary)
            if s_pre:
                _md_block(doc, s_pre.splitlines())
            for heading, body in s_sections:
                _render_v2_subheading(doc, heading)
                _md_block(doc, body)
        else:
            p = doc.add_paragraph()
            _v2_run(p, "(No summary generated for this filing.)",
                    size=brand.SZ_BODY_SM, color=_GRAY_MID, italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{number}-state-of-play-{date.today().isoformat()}.docx"
    return buf.read(), filename
